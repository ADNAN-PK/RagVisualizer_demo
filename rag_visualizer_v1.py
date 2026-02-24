# Rag_Visualizer
# to install :  pip install fastapi uvicorn scikit-learn numpy python-multipart pypdf
import uvicorn
from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.responses import HTMLResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional
import numpy as np
import re
import io
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.decomposition import PCA

# --- Configuration & App Setup ---
app = FastAPI(title="EduRAG: Interactive RAG Visualizer")

# Allow CORS for development flexibility
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- In-Memory State (Simulating a Vector DB) ---
class RAGState:
    def __init__(self):
        self.raw_text = ""
        self.chunks = []
        self.vectors = None
        self.vectorizer = None
        self.pca_2d = None
        self.coords = [] # 2D coordinates for visualization

state = RAGState()

# --- File Processing Helpers ---
def extract_text_from_file(filename: str, content: bytes) -> str:
    """
    Multi-modal text extractor. 
    Handles PDF, DOCX, TXT, MD, CSV.
    """
    filename = filename.lower()
    
    # 1. PDF Support
    if filename.endswith(".pdf"):
        try:
            import pypdf
            pdf_file = io.BytesIO(content)
            reader = pypdf.PdfReader(pdf_file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
        except ImportError:
            return "Error: Please run 'pip install pypdf' to process PDF files."
        except Exception as e:
            return f"Error reading PDF: {str(e)}"

    # 2. Word Document Support
    elif filename.endswith(".docx"):
        try:
            import docx
            doc_file = io.BytesIO(content)
            doc = docx.Document(doc_file)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text
        except ImportError:
            return "Error: Please run 'pip install python-docx' to process DOCX files."
        except Exception as e:
             return f"Error reading DOCX: {str(e)}"

    # 3. Plain Text / Markdown / Code
    else:
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            return "Error: Unsupported binary file format. Please upload PDF, DOCX, or Text files."

# --- Core RAG Logic (The "Brain") ---

def recursive_character_text_splitter(text, chunk_size, overlap):
    """
    Splits text into chunks with overlap. 
    A fundamental step in RAG to ensure context is preserved.
    """
    if not text:
        return []
    
    chunks = []
    start = 0
    text_len = len(text)
    
    while start < text_len:
        end = min(start + chunk_size, text_len)
        chunk = text[start:end]
        chunks.append(chunk)
        start += (chunk_size - overlap)
        
    return chunks

def update_knowledge_base(text: str, chunk_size: int = 200, overlap: int = 50):
    """
    Indexes the document:
    1. Chunks the text.
    2. Vectorizes chunks (Embeddings).
    3. Calculates 2D projection for visualization.
    """
    state.raw_text = text
    state.chunks = recursive_character_text_splitter(text, chunk_size, overlap)
    
    if not state.chunks:
        return {"status": "empty"}

    # 1. Vectorization (Using TF-IDF for educational clarity & speed without API keys)
    # In production, this would be OpenAI/HuggingFace embeddings.
    state.vectorizer = TfidfVectorizer(stop_words='english')
    state.vectors = state.vectorizer.fit_transform(state.chunks).toarray()
    
    # 2. Dimensionality Reduction (For Visualization)
    # We project high-dimensional vectors to 2D to show them on a chart.
    if len(state.chunks) > 2:
        pca = PCA(n_components=2)
        state.coords = pca.fit_transform(state.vectors).tolist()
        state.pca_2d = pca
    else:
        # Fallback for too few chunks
        state.coords = [[0.1, 0.1] for _ in state.chunks]
        state.pca_2d = None

    return {
        "chunk_count": len(state.chunks),
        "coords": state.coords,
        "chunks": state.chunks
    }

def retrieve_and_generate(query: str, top_k: int = 3):
    """
    The RAG Retrieval Step:
    1. Vectorize the query.
    2. Calculate Similarity (Cosine).
    3. Retrieve top K chunks.
    """
    if not state.vectorizer or state.vectors is None:
        raise HTTPException(status_code=400, detail="No document indexed.")

    # 1. Embed Query
    query_vec = state.vectorizer.transform([query]).toarray()
    
    # 2. Calculate Similarity (The "Search")
    similarities = cosine_similarity(query_vec, state.vectors).flatten()
    
    # 3. Get Top K indices
    top_indices = similarities.argsort()[-top_k:][::-1]
    
    results = []
    for idx in top_indices:
        results.append({
            "index": int(idx),
            "score": float(similarities[idx]),
            "text": state.chunks[idx]
        })

    # 4. Project Query to 2D for visualization
    query_coords = [0, 0]
    if state.pca_2d:
        query_coords = state.pca_2d.transform(query_vec)[0].tolist()

    # 5. Simulated Generation (In a real app, this sends context + query to LLM)
    # We construct a prompt to show what WOULD be sent to the LLM.
    context_str = "\n".join([f"- {r['text']}" for r in results])
    simulated_answer = (
        f"Based on the analysis of your document, here is the relevant information:\n\n"
        f"{context_str}\n\n"
        f"(Note: In a live production RAG, an LLM like GPT-4 would synthesize these chunks into a smooth answer.)"
    )

    return {
        "answer": simulated_answer,
        "retrieved_chunks": results,
        "query_coords": query_coords,
        "query_similarity_scores": similarities.tolist() # Send all scores for heatmap
    }

# --- Pydantic Models ---
class ChatRequest(BaseModel):
    query: str
    top_k: int = 3
    chunk_size: int = 300
    overlap: int = 50

# --- API Endpoints ---

@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    content = await file.read()
    
    # Extract text based on file type
    text = extract_text_from_file(file.filename, content)
    
    if text.startswith("Error"):
        return JSONResponse(status_code=400, content={"message": text, "error": True})

    # Initial default indexing
    data = update_knowledge_base(text)
    return {"message": f"Successfully processed {file.filename}", "data": data, "error": False}

@app.post("/api/reindex")
async def reindex(request: ChatRequest):
    # Allow user to change chunking params dynamically
    if not state.raw_text:
        return {"error": "No text uploaded yet"}
    data = update_knowledge_base(state.raw_text, request.chunk_size, request.overlap)
    return {"message": "Re-indexed", "data": data}

@app.post("/api/chat")
async def chat(request: ChatRequest):
    response = retrieve_and_generate(request.query, request.top_k)
    return response

# --- Frontend Application (Single Page HTML/JS/CSS) ---
# We serve this directly from the Python backend for a single-file solution.

html_content = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>EduRAG: Interactive RAG Architecture</title>
    <script src="https://cdn.tailwindcss.com"></script>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/Chart.js/3.9.1/chart.min.js"></script>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/marked/4.0.2/marked.min.js"></script>
    <link href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.0.0/css/all.min.css" rel="stylesheet">
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;600;700&display=swap');
        body { font-family: 'Inter', sans-serif; }
        .chunk-highlight { background-color: rgba(253, 224, 71, 0.3); transition: background-color 0.3s; cursor: pointer; }
        .chunk-active { background-color: rgba(253, 224, 71, 0.9); border-bottom: 2px solid #eab308; }
        .glass-panel { background: rgba(255, 255, 255, 0.9); backdrop-filter: blur(10px); border: 1px solid rgba(229, 231, 235, 0.5); }
        /* Custom Scrollbar */
        ::-webkit-scrollbar { width: 8px; }
        ::-webkit-scrollbar-track { background: #f1f1f1; }
        ::-webkit-scrollbar-thumb { background: #c7c7c7; border-radius: 4px; }
        ::-webkit-scrollbar-thumb:hover { background: #a8a8a8; }
    </style>
</head>
<body class="bg-slate-50 text-slate-800 h-screen overflow-hidden flex flex-col">

    <!-- Header -->
    <header class="bg-white border-b border-slate-200 h-14 flex items-center justify-between px-6 shrink-0 z-10 shadow-sm">
        <div class="flex items-center gap-2">
            <div class="bg-indigo-600 text-white p-1.5 rounded-lg">
                <i class="fa-solid fa-layer-group"></i>
            </div>
            <h1 class="font-bold text-lg tracking-tight text-slate-800">EduRAG <span class="text-xs font-normal text-slate-500 bg-slate-100 px-2 py-0.5 rounded-full">Interactive Architecture</span></h1>
        </div>
        <div class="flex gap-4 text-sm">
            <button onclick="document.getElementById('fileInput').click()" class="text-indigo-600 hover:text-indigo-800 font-medium transition"><i class="fa-solid fa-upload mr-1"></i> Upload File</button>
            <input type="file" id="fileInput" class="hidden" accept=".txt,.md,.pdf,.docx,.csv">
        </div>
    </header>

    <!-- Main Content -->
    <div class="flex flex-1 overflow-hidden">
        
        <!-- Left: Document Viewer (The "Corpus") -->
        <div class="w-1/3 border-r border-slate-200 flex flex-col bg-white">
            <div class="p-3 border-b border-slate-100 bg-slate-50 flex justify-between items-center">
                <h2 class="font-semibold text-xs uppercase tracking-wider text-slate-500">Source Document</h2>
                <div class="text-xs text-slate-400" id="docStats">No document loaded</div>
            </div>
            <div id="documentView" class="flex-1 overflow-y-auto p-6 text-sm leading-relaxed whitespace-pre-wrap font-mono text-slate-600">
                <div class="flex flex-col items-center justify-center h-full text-slate-400 gap-2">
                    <i class="fa-regular fa-file-lines text-4xl mb-2"></i>
                    <p class="text-center">Upload a file to see RAG in action.<br><span class="text-xs opacity-75">(PDF, DOCX, TXT, MD supported)</span></p>
                </div>
            </div>
        </div>

        <!-- Middle: RAG Internals Visualization (The "Black Box" Revealed) -->
        <div class="w-1/3 border-r border-slate-200 flex flex-col bg-slate-50 relative">
            
            <!-- Controls Overlay -->
            <div class="absolute top-4 left-4 right-4 z-10">
                <div class="glass-panel rounded-xl p-3 shadow-sm">
                    <div class="flex justify-between items-center cursor-pointer" onclick="toggleConfig()">
                        <h2 class="font-semibold text-xs uppercase tracking-wider text-indigo-600">RAG Configuration</h2>
                        <i class="fa-solid fa-gear text-slate-400 hover:text-indigo-600"></i>
                    </div>
                    <div id="configPanel" class="mt-3 space-y-3 hidden">
                        <div>
                            <label class="text-xs font-medium text-slate-600 block mb-1">Chunk Size (Chars)</label>
                            <input type="range" id="chunkSize" min="50" max="1000" step="50" value="300" class="w-full accent-indigo-600" onchange="reindexDocument()">
                            <div class="text-xs text-right text-slate-400" id="chunkSizeVal">300</div>
                        </div>
                        <div>
                            <label class="text-xs font-medium text-slate-600 block mb-1">Top-K Retrieval</label>
                            <input type="number" id="topK" value="3" min="1" max="10" class="w-full border rounded px-2 py-1 text-xs">
                        </div>
                    </div>
                </div>
            </div>

            <!-- Vector Space Chart -->
            <div class="flex-1 p-4 flex flex-col justify-center">
                <div class="relative w-full h-64 md:h-80 bg-white rounded-2xl shadow-sm border border-slate-200 p-2">
                    <canvas id="vectorChart"></canvas>
                    <div class="absolute bottom-2 right-2 text-[10px] text-slate-400 bg-white px-2 py-1 rounded border shadow-sm">
                        Visualizing Semantic Vector Space (PCA Reduced)
                    </div>
                </div>
                
                <div class="mt-4 px-2">
                    <h3 class="text-sm font-semibold text-slate-700 mb-2"><i class="fa-solid fa-microchip mr-2 text-indigo-500"></i>Process Flow</h3>
                    <div class="space-y-2 text-xs">
                        <div class="flex items-center gap-2 p-2 rounded bg-white border border-slate-100 shadow-sm transition-all" id="step1">
                            <span class="bg-slate-100 text-slate-500 w-5 h-5 flex items-center justify-center rounded-full font-bold">1</span>
                            <span>Document Chunking & Indexing</span>
                        </div>
                        <div class="flex items-center gap-2 p-2 rounded bg-white border border-slate-100 shadow-sm transition-all opacity-50" id="step2">
                            <span class="bg-slate-100 text-slate-500 w-5 h-5 flex items-center justify-center rounded-full font-bold">2</span>
                            <span>Query Vectorization</span>
                        </div>
                        <div class="flex items-center gap-2 p-2 rounded bg-white border border-slate-100 shadow-sm transition-all opacity-50" id="step3">
                            <span class="bg-slate-100 text-slate-500 w-5 h-5 flex items-center justify-center rounded-full font-bold">3</span>
                            <span>Nearest Neighbor Search (Cosine Sim)</span>
                        </div>
                    </div>
                </div>
            </div>
        </div>

        <!-- Right: Chat & Response (The "Interface") -->
        <div class="w-1/3 flex flex-col bg-white">
            <div class="p-3 border-b border-slate-100 bg-slate-50">
                <h2 class="font-semibold text-xs uppercase tracking-wider text-slate-500">AI Chat</h2>
            </div>
            
            <div id="chatHistory" class="flex-1 overflow-y-auto p-4 space-y-4">
                <div class="flex gap-3">
                    <div class="w-8 h-8 rounded-full bg-indigo-100 flex items-center justify-center text-indigo-600 shrink-0"><i class="fa-solid fa-robot"></i></div>
                    <div class="bg-slate-100 p-3 rounded-2xl rounded-tl-none text-sm text-slate-700">
                        Hello! Upload a document (PDF, Word, or Text), and I'll show you exactly how I read it to answer your questions.
                    </div>
                </div>
            </div>

            <div class="p-4 border-t border-slate-100">
                <form id="chatForm" class="relative">
                    <input type="text" id="userQuery" placeholder="Ask about the document..." 
                           class="w-full bg-slate-50 border border-slate-200 rounded-xl py-3 pl-4 pr-12 text-sm focus:outline-none focus:ring-2 focus:ring-indigo-500 transition shadow-sm">
                    <button type="submit" class="absolute right-2 top-2 p-1.5 bg-indigo-600 text-white rounded-lg hover:bg-indigo-700 transition">
                        <i class="fa-solid fa-paper-plane text-xs"></i>
                    </button>
                </form>
            </div>
        </div>
    </div>

    <script>
        // --- Global State ---
        let chartInstance = null;
        let documentChunks = [];
        let chunkCoords = [];

        // --- Event Listeners ---
        document.getElementById('chunkSize').addEventListener('input', (e) => {
            document.getElementById('chunkSizeVal').textContent = e.target.value;
        });

        document.getElementById('fileInput').addEventListener('change', async (e) => {
            const file = e.target.files[0];
            if (!file) return;

            const formData = new FormData();
            formData.append('file', file);

            setLoading(true);
            try {
                const res = await fetch('/api/upload', { method: 'POST', body: formData });
                const data = await res.json();
                
                if (data.error) {
                    alert(data.message);
                    return;
                }

                handleIndexUpdate(data.data);
                addMessage("system", `Processed ${file.name}. Split into ${data.data.chunk_count} chunks.`);
            } catch (err) {
                alert("Error uploading file");
                console.error(err);
            } finally {
                setLoading(false);
            }
        });

        document.getElementById('chatForm').addEventListener('submit', async (e) => {
            e.preventDefault();
            const query = document.getElementById('userQuery').value;
            if (!query) return;

            // UI Updates
            addMessage("user", query);
            document.getElementById('userQuery').value = '';
            highlightProcess(true);

            // API Call
            try {
                const res = await fetch('/api/chat', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({
                        query: query,
                        top_k: parseInt(document.getElementById('topK').value),
                        chunk_size: parseInt(document.getElementById('chunkSize').value),
                        overlap: 50
                    })
                });
                const data = await res.json();

                // Visualize Retrieval
                updateChartWithQuery(data.query_coords, data.retrieved_chunks);
                highlightChunksInText(data.retrieved_chunks);
                addMessage("ai", data.answer);
                
                // Reset Process Highlight
                setTimeout(() => highlightProcess(false), 2000);

            } catch (err) {
                addMessage("system", "Error getting answer. Make sure a document is uploaded.");
            }
        });

        async function reindexDocument() {
            setLoading(true);
            try {
                const res = await fetch('/api/reindex', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({
                        query: "",
                        chunk_size: parseInt(document.getElementById('chunkSize').value),
                        overlap: 50
                    })
                });
                const data = await res.json();
                handleIndexUpdate(data.data);
            } catch (err) {
                console.error(err);
            } finally {
                setLoading(false);
            }
        }

        // --- Core Functions ---

        function handleIndexUpdate(data) {
            documentChunks = data.chunks;
            chunkCoords = data.coords;
            
            // Update Text View
            const docView = document.getElementById('documentView');
            docView.innerHTML = '';
            documentChunks.forEach((chunk, idx) => {
                const span = document.createElement('span');
                span.textContent = chunk;
                span.id = `chunk-${idx}`;
                span.className = 'chunk-highlight border-r border-dotted border-slate-300 pr-1 mr-1 text-slate-500 hover:text-slate-800';
                span.title = `Chunk ID: ${idx}`;
                span.onclick = () => highlightSingleChunk(idx);
                docView.appendChild(span);
            });

            document.getElementById('docStats').textContent = `${data.chunk_count} chunks`;
            
            // Init Chart
            initChart(data.coords);
        }

        function highlightChunksInText(retrieved) {
            // Reset all
            document.querySelectorAll('.chunk-highlight').forEach(el => {
                el.classList.remove('chunk-active', 'bg-yellow-100', 'text-slate-900');
                el.classList.add('text-slate-500');
            });

            // Highlight retrieved
            retrieved.forEach(item => {
                const el = document.getElementById(`chunk-${item.index}`);
                if (el) {
                    el.classList.add('chunk-active', 'text-slate-900');
                    el.classList.remove('text-slate-500');
                    el.scrollIntoView({ behavior: 'smooth', block: 'center' });
                }
            });
        }
        
        function highlightSingleChunk(idx) {
             document.querySelectorAll('.chunk-highlight').forEach(el => el.classList.remove('chunk-active'));
             const el = document.getElementById(`chunk-${idx}`);
             if(el) el.classList.add('chunk-active');
        }

        function addMessage(role, text) {
            const history = document.getElementById('chatHistory');
            const div = document.createElement('div');
            div.className = "flex gap-3 animated fade-in-up";
            
            const isUser = role === "user";
            const icon = isUser ? '<i class="fa-solid fa-user"></i>' : (role === 'system' ? '<i class="fa-solid fa-info"></i>' : '<i class="fa-solid fa-robot"></i>');
            const bgClass = isUser ? 'bg-indigo-600 text-white rounded-tr-none' : (role === 'system' ? 'bg-red-50 text-red-600 border border-red-100' : 'bg-slate-100 text-slate-700 rounded-tl-none');
            const iconBg = isUser ? 'bg-indigo-100 text-indigo-600' : 'bg-indigo-50 text-indigo-600';

            div.innerHTML = `
                <div class="w-8 h-8 rounded-full ${isUser ? 'order-2 bg-slate-200 text-slate-600' : iconBg} flex items-center justify-center shrink-0 text-xs">
                    ${icon}
                </div>
                <div class="${bgClass} p-3 rounded-2xl text-sm shadow-sm max-w-[85%]">
                    ${role === 'ai' ? marked.parse(text) : text}
                </div>
            `;
            history.appendChild(div);
            history.scrollTop = history.scrollHeight;
        }

        // --- Visualization Logic (Chart.js) ---

        function initChart(coords) {
            const ctx = document.getElementById('vectorChart').getContext('2d');
            
            if (chartInstance) chartInstance.destroy();

            const dataPoints = coords.map((c, i) => ({ x: c[0], y: c[1], chunkId: i }));

            chartInstance = new Chart(ctx, {
                type: 'scatter',
                data: {
                    datasets: [{
                        label: 'Knowledge Chunks',
                        data: dataPoints,
                        backgroundColor: '#94a3b8',
                        pointRadius: 4,
                        pointHoverRadius: 6
                    }, {
                        label: 'Query Vector',
                        data: [],
                        backgroundColor: '#ef4444',
                        pointRadius: 8,
                        pointStyle: 'star'
                    }, {
                        label: 'Retrieved',
                        data: [],
                        backgroundColor: '#eab308',
                        pointRadius: 6,
                        pointStyle: 'circle'
                    }]
                },
                options: {
                    responsive: true,
                    maintainAspectRatio: false,
                    plugins: {
                        legend: { position: 'bottom', labels: { usePointStyle: true, font: { size: 10 } } },
                        tooltip: {
                            callbacks: {
                                label: (ctx) => {
                                    const idx = ctx.raw.chunkId;
                                    if (idx !== undefined) return `Chunk ${idx}: ${documentChunks[idx].substring(0, 30)}...`;
                                    return ctx.dataset.label;
                                }
                            }
                        }
                    },
                    scales: {
                        x: { display: false },
                        y: { display: false }
                    }
                }
            });
        }

        function updateChartWithQuery(queryCoords, retrieved) {
            if (!chartInstance) return;

            // Update Query Point
            chartInstance.data.datasets[1].data = [{ x: queryCoords[0], y: queryCoords[1] }];

            // Update Retrieved Points (Find coords for retrieved indices)
            const retrievedPoints = retrieved.map(r => {
                const c = chunkCoords[r.index];
                return { x: c[0], y: c[1], chunkId: r.index };
            });
            chartInstance.data.datasets[2].data = retrievedPoints;

            chartInstance.update();
        }

        // --- Helpers ---
        function toggleConfig() {
            document.getElementById('configPanel').classList.toggle('hidden');
        }
        
        function setLoading(isLoading) {
             const btn = document.querySelector('button[type="submit"]');
             if(isLoading) {
                 btn.innerHTML = '<i class="fa-solid fa-spinner fa-spin"></i>';
                 btn.disabled = true;
             } else {
                 btn.innerHTML = '<i class="fa-solid fa-paper-plane"></i>';
                 btn.disabled = false;
             }
        }

        function highlightProcess(active) {
            const steps = ['step1', 'step2', 'step3'];
            let delay = 0;
            if (active) {
                steps.forEach((id, i) => {
                    setTimeout(() => {
                        document.getElementById(id).classList.remove('opacity-50');
                        document.getElementById(id).classList.add('bg-indigo-50', 'border-indigo-200');
                    }, delay);
                    delay += 300;
                });
            } else {
                steps.forEach(id => {
                    document.getElementById(id).classList.add('opacity-50');
                    document.getElementById(id).classList.remove('bg-indigo-50', 'border-indigo-200');
                });
            }
        }
    </script>
</body>
</html>
"""

@app.get("/", response_class=HTMLResponse)
async def get_app():
    return html_content

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)